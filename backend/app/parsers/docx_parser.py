"""
DOCX text extraction using python-docx.

Extracts all paragraphs and table content with structure preservation.
"""
import time
from docx import Document
from typing import List, Tuple
from app.parsers.pdf_parser import ExtractionResult


def extract_text_from_docx(file_path: str) -> ExtractionResult:
    """
    Extract text from a DOCX file preserving paragraph and table structure.
    """
    start = time.time()

    try:
        doc = Document(file_path)
    except Exception as e:
        elapsed = (time.time() - start) * 1000
        return ExtractionResult(
            text="",
            page_count=0,
            extraction_method="docx",
            quality_score=0.0,
            quality_issues=[f"DOCX_EXTRACTION_FAILED: {str(e)}"],
            extraction_time_ms=elapsed,
        )

    sections: List[str] = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            sections.append(text)

    # Extract tables
    for table_idx, table in enumerate(doc.tables):
        table_lines: List[str] = []
        for row in table.rows:
            cells_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    cells_text.append(cell_text)
            if cells_text:
                table_lines.append(" | ".join(cells_text))

        if table_lines:
            sections.append("[TABLE]")
            sections.extend(table_lines)
            sections.append("[/TABLE]")

    full_text = "\n".join(sections)
    elapsed = (time.time() - start) * 1000

    quality_score, quality_issues = _assess_docx_quality(full_text)
    lines = [l for l in full_text.split("\n") if l.strip()]

    return ExtractionResult(
        text=full_text,
        page_count=0,  # DOCX doesn't have explicit pages
        extraction_method="docx",
        ocr_used=False,
        quality_score=quality_score,
        quality_issues=quality_issues,
        extraction_time_ms=elapsed,
        line_count=len(lines),
        char_count=len(full_text),
    )


def _assess_docx_quality(text: str) -> Tuple[float, List[str]]:
    """Assess quality of DOCX extraction."""
    issues: List[str] = []

    if not text or not text.strip():
        return 0.0, ["TEXT_EMPTY"]

    lines = [l for l in text.split("\n") if l.strip()]
    if len(lines) < 3:
        issues.append("TOO_FEW_LINES")

    if len(text.strip()) < 100:
        issues.append("VERY_SHORT_CONTENT")

    score = 1.0
    if "TOO_FEW_LINES" in issues:
        score -= 0.3
    if "VERY_SHORT_CONTENT" in issues:
        score -= 0.3

    return max(0.0, score), issues
